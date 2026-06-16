"""Report Writer Agent — v3 (section-based, professional output).

Key differences vs v2:
  * The report is assembled SECTION BY SECTION. Each section places its own
    table(s) and chart(s) right next to the AI-written analysis for that
    section — nothing is dumped at the end of the report.
  * Charts never break a sentence: they are placed on their own, between the
    table and the prose, at section boundaries only.
  * HTML adopts a polished, professional template (summary box, section
    accents, chart captions) modeled on a hand-made analyst report.
  * PDF generation has a 3-tier fallback so it works in almost any
    environment: WeasyPrint -> wkhtmltopdf (pdfkit) -> reportlab.
  * DOCX is built from the same structured section model, so its layout
    matches the HTML.

The report is represented internally as a list of "blocks". Each block is one
of: heading, paragraph, table, chart, caption, summary, rule, meta. Every
output format (Markdown / HTML / DOCX) is rendered from this single block list,
which guarantees the three formats stay consistent.
"""
import os
import re
import base64
from datetime import datetime

import markdown as md_lib
from state import AnalysisState
from utils.ratios import RATIO_LABELS, RATIO_CATEGORIES, PERCENT_RATIOS


# =========================================================================
# Block model
# =========================================================================
# A block is a dict with a "type" key. Helpers below build them.

def _h(level, text):           return {"type": "heading", "level": level, "text": text}
def _p(text):                  return {"type": "paragraph", "text": text}
def _summary(text):            return {"type": "summary", "text": text}
def _caption(text):            return {"type": "caption", "text": text}
def _rule():                   return {"type": "rule"}
def _meta(text):               return {"type": "meta", "text": text}
def _table(headers, rows, widths=None):
    return {"type": "table", "headers": headers, "rows": rows, "widths": widths}
def _chart(path, caption=""):  return {"type": "chart", "path": path, "caption": caption}


def _fmt(val, as_pct=False):
    if val is None:
        return "—"
    if as_pct:
        return f"{val * 100:.1f}%"
    return f"{val:.2f}"


# =========================================================================
# Section builders — each returns a list of blocks (heading + data + AI text)
# =========================================================================

def _section_text(sections: dict, key: str, fallback: str = "") -> list:
    """Return the AI-written prose for a section id as paragraph blocks."""
    text = sections.get(key, "").strip()
    if not text:
        return [_p(fallback)] if fallback else []
    # Split prose into paragraphs; strip any stray markdown heading lines.
    blocks = []
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        # drop accidental '##' headings the model may emit
        para = re.sub(r"^#{1,6}\s*", "", para)
        blocks.append(_p(para))
    return blocks


def _ratio_table(by_year: dict) -> dict:
    years = sorted(by_year.keys(), reverse=True)
    headers = ["Ratio"] + [str(y) for y in years]
    rows = []
    for category, keys in RATIO_CATEGORIES.items():
        rows.append({"_category": category})  # marker row
        for key in keys:
            label = RATIO_LABELS[key]
            as_pct = key in PERCENT_RATIOS
            rows.append([label] + [_fmt(by_year[y].get(key), as_pct) for y in years])
    return _table(headers, rows)


def _dupont_table(by_year: dict) -> dict:
    years = sorted(by_year.keys(), reverse=True)
    components = [
        ("Net Margin", "net_margin", True),
        ("Asset Turnover", "asset_turnover", False),
        ("Equity Multiplier", "equity_multiplier", False),
        ("→ ROE (3-step)", "roe_3step", True),
        ("Tax Burden", "tax_burden", False),
        ("Interest Burden", "interest_burden", False),
        ("Operating Margin", "operating_margin", True),
        ("→ ROE (5-step)", "roe_5step", True),
    ]
    headers = ["Component"] + [str(y) for y in years]
    rows = []
    for label, key, as_pct in components:
        cells = [label]
        for y in years:
            d = by_year[y].get("_dupont", {})
            cells.append(_fmt(d.get(key), as_pct))
        rows.append(cells)
    return _table(headers, rows)


def _trend_table(ticker_trends: dict) -> dict | None:
    headers = ["Ratio", "Direction", "CAGR", "First", "Last"]
    rows = []
    for ratio_key, t in ticker_trends.items():
        if t["direction"] == "insufficient_data":
            continue
        label = RATIO_LABELS.get(ratio_key, ratio_key)
        cagr = f"{t['cagr']*100:.1f}%" if t["cagr"] is not None else "—"
        as_pct = ratio_key in PERCENT_RATIOS
        first = (f"{t['first_year']}: {_fmt(t['first_value'], as_pct)}"
                 if t["first_value"] is not None else "—")
        last = (f"{t['last_year']}: {_fmt(t['last_value'], as_pct)}"
                if t["last_value"] is not None else "—")
        rows.append([label, t["direction"].replace("_", " "), cagr, first, last])
    return _table(headers, rows) if rows else None


def _peer_table(peer_for_primary: dict) -> dict | None:
    headers = ["Ratio", "Company", "Peer Avg", "Peer Median", "Percentile"]
    rows = []
    for ratio_key, d in peer_for_primary.items():
        label = RATIO_LABELS.get(ratio_key, ratio_key)
        as_pct = ratio_key in PERCENT_RATIOS
        rows.append([
            label,
            _fmt(d["primary_value"], as_pct),
            _fmt(d["peer_average"], as_pct),
            _fmt(d["peer_median"], as_pct),
            f"{d['percentile_rank']:.0f}th",
        ])
    return _table(headers, rows) if rows else None


def _fmt_money(val) -> str:
    if val is None:
        return "—"
    a = abs(val)
    if a >= 1e12:
        return f"${val/1e12:.2f}T"
    if a >= 1e9:
        return f"${val/1e9:.1f}B"
    if a >= 1e6:
        return f"${val/1e6:.1f}M"
    return f"${val:,.0f}"


def _fmt_x(val) -> str:
    return "—" if val is None else f"{val:.1f}×"


def _growth_table(growth: dict) -> dict | None:
    rows = []
    for key in ["revenue", "net_income", "eps", "fcf"]:
        g = growth.get(key)
        if not g:
            continue
        is_eps = key == "eps"
        first = f"${g['first']:.2f}" if is_eps else _fmt_money(g["first"])
        last = f"${g['last']:.2f}" if is_eps else _fmt_money(g["last"])
        cagr = f"{g['cagr']*100:.1f}%" if g.get("cagr") is not None else "—"
        rows.append([g["label"], f"{g['first_year']}: {first}",
                     f"{g['last_year']}: {last}", cagr])
    return _table(["Metric", "First", "Latest", "CAGR"], rows) if rows else None


def _valuation_table(valuation: dict, val_peers: dict) -> dict | None:
    defs = [
        ("pe", "P / E", "x"), ("ps", "P / S", "x"), ("pb", "P / B", "x"),
        ("ev_ebitda", "EV / EBITDA", "x"),
        ("fcf_yield", "FCF Yield", "pct"),
        ("earnings_yield", "Earnings Yield", "pct"),
        ("dividend_yield", "Dividend Yield", "pct"),
    ]
    has_peers = bool(val_peers)
    headers = ["Metric", "Company"] + (["Peer Avg", "Peer Median", "Percentile"]
                                       if has_peers else [])
    rows = []
    for key, label, kind in defs:
        v = valuation.get(key)
        d = val_peers.get(key) if has_peers else None
        if v is None and not d:
            continue
        comp = _fmt_x(v) if kind == "x" else _fmt(v, as_pct=True)
        row = [label, comp]
        if has_peers:
            if d:
                pa = _fmt_x(d["peer_average"]) if kind == "x" else _fmt(d["peer_average"], as_pct=True)
                pm = _fmt_x(d["peer_median"]) if kind == "x" else _fmt(d["peer_median"], as_pct=True)
                row += [pa, pm, f"{d['percentile_rank']:.0f}th"]
            else:
                row += ["—", "—", "—"]
        rows.append(row)
    return _table(headers, rows) if rows else None


def _capalloc_table(cap: dict) -> dict | None:
    def pct(k):
        v = cap.get(k)
        return "—" if v is None else f"{v*100:.1f}%"

    rows = [
        ["Buybacks (latest FY)", _fmt_money(cap.get("buybacks"))],
        ["Dividends paid (latest FY)", _fmt_money(cap.get("dividends"))],
        ["Total returned to shareholders", _fmt_money(cap.get("total_returned"))],
        ["Dividend payout ratio", pct("payout_ratio")],
        ["Total payout ratio (div + buyback)", pct("total_payout_ratio")],
        ["Shareholder yield", pct("shareholder_yield")],
    ]
    if cap.get("share_count_change") is not None:
        rows.append(["Share count change (period)", pct("share_count_change")])
    return _table(["Capital allocation", "Value"], rows)


def _comparison_table(comparison: dict) -> dict | None:
    rankings = comparison.get("rankings", {})
    snapshots = comparison.get("snapshots", {})
    tickers = list(snapshots.keys())
    if not tickers or not rankings:
        return None
    headers = ["Ratio"] + tickers + ["Best"]
    rows = []
    for key, r in rankings.items():
        as_pct = key in PERCENT_RATIOS
        cells = [r["label"]] + [_fmt(r["values"].get(t), as_pct) for t in tickers]
        cells.append(r["best"])
        rows.append(cells)
    return _table(headers, rows)


def _build_blocks(state: AnalysisState) -> list:
    """Assemble the full ordered block list for the report."""
    tickers = state.get("tickers", [])
    raw = state.get("raw_data", {})
    quality = state.get("data_quality", {})
    ratios = state.get("ratios", {})
    trends = state.get("trends", {})
    peers_map = state.get("peers", {})
    peer_analysis = state.get("peer_analysis", {})
    fundamentals = state.get("fundamentals", {})
    valuation_peers = state.get("valuation_peers", {})
    comparison = state.get("comparison", {})
    sections = state.get("section_analysis", {})
    critique = state.get("critique", "")
    query = state.get("user_query", "")
    charts = state.get("charts", {})
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    blocks: list = []

    # ---- Title + meta ----
    if len(tickers) == 1:
        info = raw.get(tickers[0], {}).get("info", {})
        name = info.get("longName", tickers[0])
        blocks.append(_h(1, f"{name} ({tickers[0]})"))
    else:
        blocks.append(_h(1, "Financial Statement Analysis Report"))

    meta_bits = [f"Generated {now}"]
    if len(tickers) == 1:
        info = raw.get(tickers[0], {}).get("info", {})
        sector = info.get("sector", "Unknown")
        industry = info.get("industry", "Unknown")
        mcap = info.get("marketCap")
        meta_bits.append(f"Sector: {sector} / {industry}")
        if mcap:
            meta_bits.append(f"Market cap: ${mcap/1e9:.1f}B")
    blocks.append(_meta("  |  ".join(meta_bits)))
    blocks.append(_meta(f"Original request: {query}"))

    # ---- Intro summary box ----
    blocks.append(_summary(
        "This report was produced by a multi-agent AI workflow. Each section "
        "first presents the underlying data (tables and charts), then provides "
        "an AI-generated analysis grounded in those specific numbers and "
        "visuals. Ratios and the DuPont decomposition are computed "
        "deterministically; the narrative is written by an AI analyst and "
        "reviewed by an automated critic."
    ))

    # ---- Executive Summary ----
    blocks.append(_h(2, "Executive Summary"))
    blocks += _section_text(sections, "executive_summary",
                            "_No executive summary was generated._")

    # ---- Companies Analyzed ----
    blocks.append(_h(2, "Companies Analyzed"))
    comp_headers = ["Ticker", "Company", "Sector", "Industry", "Market Cap"]
    comp_rows = []
    for t in tickers:
        info = raw.get(t, {}).get("info", {})
        mcap = info.get("marketCap")
        comp_rows.append([
            t,
            info.get("longName", t),
            info.get("sector", "Unknown"),
            info.get("industry", "Unknown"),
            f"${mcap/1e9:.1f}B" if mcap else "—",
        ])
    blocks.append(_table(comp_headers, comp_rows))

    if peers_map:
        peer_lines = []
        for primary, peer_list in peers_map.items():
            if peer_list:
                peer_lines.append(f"{primary}: {', '.join(peer_list)}")
        if peer_lines:
            blocks.append(_p("AI-selected peer companies — " + "; ".join(peer_lines) + "."))

    # ---- Data Quality ----
    blocks.append(_h(2, "Data Quality Review"))
    dq_rows = []
    for t, q in quality.items():
        issues = "; ".join(q.get("issues", [])) or "None"
        dq_rows.append([
            t,
            f"{q['completeness']*100:.0f}%",
            "Pass" if q["status"] == "pass" else "Fail",
            issues,
        ])
    blocks.append(_table(["Ticker", "Completeness", "Status", "Issues"], dq_rows))

    # ---- Per-company analytical sections ----
    for t in tickers:
        if t not in ratios:
            continue
        by_year = ratios[t]

        blocks.append(_rule())
        blocks.append(_h(2, f"{t} — Detailed Analysis"))

        # Full ratio table for the company
        blocks.append(_h(3, f"{t}: Financial Ratios"))
        blocks.append(_ratio_table(by_year))

        fund = fundamentals.get(t, {})

        # Growth & Scale (absolute figures + growth rates)
        if fund.get("growth"):
            gt = _growth_table(fund["growth"])
            if gt:
                blocks.append(_h(3, f"{t}: Growth & Scale"))
                blocks.append(gt)
                blocks += _section_text(sections, f"growth_{t}")

        # Profitability
        blocks.append(_h(3, f"{t}: Profitability"))
        if f"profitability_{t}" in charts:
            blocks.append(_chart(charts[f"profitability_{t}"],
                                 f"{t} profitability margins over time"))
        blocks += _section_text(sections, f"profitability_{t}")

        # DuPont
        blocks.append(_h(3, f"{t}: DuPont Decomposition"))
        blocks.append(_dupont_table(by_year))
        if f"dupont_{t}" in charts:
            blocks.append(_chart(charts[f"dupont_{t}"],
                                 f"{t} DuPont drivers and resulting ROE"))
        blocks += _section_text(sections, f"dupont_{t}")

        # Trend
        if t in trends:
            tt = _trend_table(trends[t])
            if tt:
                blocks.append(_h(3, f"{t}: Trend Analysis"))
                blocks.append(tt)
                for rk in ["roe", "net_margin", "roa"]:
                    ck = f"trend_{t}_{rk}"
                    if ck in charts:
                        blocks.append(_chart(charts[ck],
                                             f"{t} {RATIO_LABELS.get(rk, rk)} trend"))
                blocks += _section_text(sections, f"trend_{t}")

        # Liquidity & Leverage
        blocks.append(_h(3, f"{t}: Liquidity & Leverage"))
        blocks += _section_text(sections, f"liquidity_leverage_{t}")

        # Valuation (multiples + peer comparison)
        if fund.get("valuation"):
            vt = _valuation_table(fund["valuation"], valuation_peers.get(t, {}))
            if vt:
                blocks.append(_h(3, f"{t}: Valuation"))
                blocks.append(vt)
                blocks += _section_text(sections, f"valuation_{t}")

        # Capital Allocation (buybacks, dividends, share count)
        if fund.get("capital_allocation"):
            ct = _capalloc_table(fund["capital_allocation"])
            if ct:
                blocks.append(_h(3, f"{t}: Capital Allocation"))
                blocks.append(ct)
                blocks += _section_text(sections, f"capital_allocation_{t}")

        # Peer benchmarking
        if t in peer_analysis:
            pt = _peer_table(peer_analysis[t])
            if pt:
                blocks.append(_h(3, f"{t}: Peer Benchmarking"))
                blocks.append(pt)
                if any(d.get("excluded_peers") for d in peer_analysis[t].values()):
                    blocks.append(_meta(
                        "Note: peers with negative book equity were excluded from "
                        "equity-based ratios (ROE, equity multiplier, debt/equity) so a "
                        "single distorted value does not poison the peer average."))
                pchart_keys = [k for k in charts if k.startswith(f"peer_{t}_")]
                for ck in pchart_keys[:3]:
                    blocks.append(_chart(charts[ck]))
                blocks += _section_text(sections, f"peer_{t}")

    # ---- Multi-company comparison ----
    comp_table = _comparison_table(comparison)
    if comp_table:
        blocks.append(_rule())
        blocks.append(_h(2, "Side-by-Side Comparison (Most Recent Year)"))
        blocks.append(comp_table)
        if "win_tally" in charts:
            blocks.append(_chart(charts["win_tally"], "Best-in-class wins per company"))
        for rk in ["roe", "net_margin", "roa", "debt_to_equity"]:
            ck = f"compare_{rk}"
            if ck in charts:
                blocks.append(_chart(charts[ck]))
        blocks += _section_text(sections, "comparison")

    # ---- Risks & Caveats ----
    blocks.append(_rule())
    blocks.append(_h(2, "Risks & Caveats"))
    blocks += _section_text(sections, "risks")

    # ---- Critic review ----
    if critique:
        blocks.append(_rule())
        blocks.append(_h(2, "Critic Review"))
        blocks.append(_p("The critic agent reviewed the analyst's draft for "
                         "accuracy and completeness before approval."))
        blocks.append(_summary(critique))

    # ---- Methodology ----
    blocks.append(_rule())
    blocks.append(_h(2, "Methodology"))
    blocks.append(_p(
        "This report was generated by a multi-agent workflow: Planner → "
        "Retriever → Peer Selector → Peer Retriever → Validator → Ratios → "
        "Trend Analyzer → Comparator → Peer Analyzer → Fundamentals "
        "(valuation, growth, capital allocation) → Chart Builder → Analyst "
        "(vision-enabled, section by section) → Critic → Report Writer. Data "
        "source: Yahoo Finance. Ratios, the DuPont decomposition, valuation "
        "multiples, growth and capital-allocation figures are computed "
        "deterministically; the per-section narrative is generated by an AI "
        "analyst from the numerical tables and chart images, then reviewed by a "
        "critic agent. Valuation multiples use latest-fiscal-year fundamentals "
        "against current market capitalization (a trailing approximation)."
    ))
    blocks.append(_meta("Not investment advice."))

    return blocks


# =========================================================================
# Renderers
# =========================================================================

def _render_table_md(tb: dict) -> str:
    headers = tb["headers"]
    out = ["| " + " | ".join(headers) + " |",
           "|" + "---|" * len(headers)]
    for row in tb["rows"]:
        if isinstance(row, dict) and "_category" in row:
            out.append(f"| **{row['_category']}** |" + " |" * (len(headers) - 1))
            continue
        cells = [str(c) for c in row]
        cells += [""] * (len(headers) - len(cells))
        out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


def _build_markdown(blocks: list) -> str:
    out = []
    for b in blocks:
        t = b["type"]
        if t == "heading":
            out.append(f"{'#' * b['level']} {b['text']}")
        elif t == "paragraph":
            out.append(b["text"])
        elif t == "meta":
            out.append(f"*{b['text']}*")
        elif t == "summary":
            out.append(f"> {b['text']}")
        elif t == "caption":
            out.append(f"*{b['text']}*")
        elif t == "rule":
            out.append("---")
        elif t == "table":
            out.append(_render_table_md(b))
        elif t == "chart":
            # Markdown keeps a stable placeholder so the Streamlit app can
            # still render charts inline; it is ALWAYS on its own line/block.
            name = os.path.splitext(os.path.basename(b["path"]))[0] if b.get("path") else ""
            out.append(f"[CHART: {name}]")
            if b.get("caption"):
                out.append(f"*{b['caption']}*")
        out.append("")  # blank line between blocks
    return "\n".join(out)


# ---------- HTML ----------

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  body {{
    font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
    max-width: 920px; margin: 32px auto; padding: 24px;
    color: #222; line-height: 1.7; background: #fff;
  }}
  h1 {{ color: #1f4e79; border-bottom: 3px solid #c00000;
        padding-bottom: 12px; margin-bottom: 6px; font-size: 1.9rem; }}
  h2 {{ color: #1f4e79; border-left: 4px solid #c00000; padding-left: 12px;
        margin-top: 38px; font-size: 1.35rem; }}
  h3 {{ color: #2c5282; margin-top: 26px; font-size: 1.1rem; }}
  p {{ margin: 12px 0; }}
  .meta {{ color: #555; font-size: 13px; margin: 4px 0; }}
  .summary-box {{ background: #f5f7fa; border-left: 4px solid #1f4e79;
                  padding: 14px 20px; margin: 20px 0; border-radius: 4px;
                  font-size: 14.5px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 16px 0;
           font-size: 13.5px; }}
  th, td {{ border: 1px solid #ddd; padding: 7px 11px; text-align: left; }}
  th {{ background: #1f4e79; color: #fff; }}
  tr:nth-child(even) td {{ background: #fafbfc; }}
  td.category {{ background: #eef2f7; font-weight: 600; color: #1f4e79; }}
  .chart-block {{ margin: 22px 0; text-align: center; }}
  .chart-block img {{ max-width: 86%; height: auto;
                      border: 1px solid #e5e7eb; border-radius: 6px; }}
  .chart-caption {{ font-size: 12.5px; color: #666; margin-top: 6px;
                    font-style: italic; }}
  hr {{ border: none; border-top: 1px solid #e6e6e6; margin: 28px 0; }}
  .footer {{ margin-top: 44px; padding-top: 14px; border-top: 1px solid #eee;
             color: #888; font-size: 12px; text-align: center; }}
  strong {{ color: #1f4e79; }}
</style>
</head>
<body>
{body}
<div class="footer">Generated by the Financial Statement Analysis Agent · Not investment advice.</div>
</body>
</html>"""


def _md_inline(text: str) -> str:
    """Render inline markdown (**bold**, *italic*) inside a paragraph to HTML."""
    html = md_lib.markdown(text, extensions=["tables"])
    # strip the wrapping <p> markdown adds, we control block tags ourselves
    html = re.sub(r"^<p>(.*)</p>$", r"\1", html.strip(), flags=re.DOTALL)
    return html


def _img_data_uri(path: str) -> str | None:
    if not path or not os.path.exists(path):
        return None
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return f"data:image/png;base64,{b64}"


def _render_table_html(tb: dict) -> str:
    headers = tb["headers"]
    ncols = len(headers)
    out = ["<table><thead><tr>"]
    out += [f"<th>{h}</th>" for h in headers]
    out.append("</tr></thead><tbody>")
    for row in tb["rows"]:
        if isinstance(row, dict) and "_category" in row:
            out.append(f'<tr><td class="category" colspan="{ncols}">'
                       f'{row["_category"]}</td></tr>')
            continue
        cells = [str(c) for c in row]
        cells += [""] * (ncols - len(cells))
        out.append("<tr>" + "".join(f"<td>{_md_inline(c)}</td>" for c in cells) + "</tr>")
    out.append("</tbody></table>")
    return "".join(out)


def _build_html(blocks: list, title: str) -> str:
    body = []
    for b in blocks:
        t = b["type"]
        if t == "heading":
            body.append(f"<h{b['level']}>{b['text']}</h{b['level']}>")
        elif t == "paragraph":
            body.append(f"<p>{_md_inline(b['text'])}</p>")
        elif t == "meta":
            body.append(f'<p class="meta">{_md_inline(b["text"])}</p>')
        elif t == "summary":
            body.append(f'<div class="summary-box">{_md_inline(b["text"])}</div>')
        elif t == "caption":
            body.append(f'<p class="chart-caption">{_md_inline(b["text"])}</p>')
        elif t == "rule":
            body.append("<hr>")
        elif t == "table":
            body.append(_render_table_html(b))
        elif t == "chart":
            uri = _img_data_uri(b.get("path", ""))
            if uri:
                cap = (f'<div class="chart-caption">{b["caption"]}</div>'
                       if b.get("caption") else "")
                body.append(f'<div class="chart-block"><img src="{uri}" '
                            f'alt="chart"/>{cap}</div>')
    return HTML_TEMPLATE.format(title=title, body="\n".join(body))


# ---------- DOCX ----------

def _build_docx(blocks: list, out_path: str) -> str:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    NAVY = RGBColor(0x1F, 0x4E, 0x79)

    def add_runs(paragraph, text):
        """Add text to a paragraph, honoring **bold** / *italic* markers."""
        for part in re.split(r"(\*\*.+?\*\*|\*.+?\*)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2]); r.bold = True
            elif part.startswith("*") and part.endswith("*"):
                r = paragraph.add_run(part[1:-1]); r.italic = True
            else:
                paragraph.add_run(part)

    for b in blocks:
        t = b["type"]
        if t == "heading":
            lvl = min(b["level"], 4)
            doc.add_heading(b["text"], level=0 if lvl == 1 else lvl)
        elif t == "paragraph":
            p = doc.add_paragraph()
            add_runs(p, b["text"])
        elif t == "meta":
            p = doc.add_paragraph()
            r = p.add_run(b["text"]); r.italic = True
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        elif t == "summary":
            p = doc.add_paragraph()
            r = p.add_run(b["text"]); r.italic = True
            p.paragraph_format.left_indent = Inches(0.3)
        elif t == "caption":
            p = doc.add_paragraph()
            r = p.add_run(b["text"]); r.italic = True
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "rule":
            p = doc.add_paragraph()
            r = p.add_run("─" * 45)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "table":
            headers = b["headers"]
            ncols = len(headers)
            display_rows = [headers]
            for row in b["rows"]:
                if isinstance(row, dict) and "_category" in row:
                    display_rows.append(["§" + row["_category"]] + [""] * (ncols - 1))
                else:
                    cells = [str(c) for c in row]
                    cells += [""] * (ncols - len(cells))
                    display_rows.append(cells)
            table = doc.add_table(rows=len(display_rows), cols=ncols)
            try:
                table.style = "Light Grid Accent 1"
            except Exception:
                pass
            for i, rdata in enumerate(display_rows):
                for j, cell_text in enumerate(rdata):
                    cell = table.cell(i, j)
                    txt = cell_text.replace("**", "")
                    is_category = txt.startswith("§")
                    if is_category:
                        txt = txt[1:]
                    cell.text = txt
                    for run in cell.paragraphs[0].runs:
                        if i == 0:
                            run.bold = True
                        if is_category:
                            run.bold = True
                            run.font.color.rgb = NAVY
            doc.add_paragraph()
        elif t == "chart":
            path = b.get("path", "")
            if path and os.path.exists(path):
                try:
                    doc.add_picture(path, width=Inches(5.6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if b.get("caption"):
                        cp = doc.add_paragraph()
                        cr = cp.add_run(b["caption"])
                        cr.italic = True; cr.font.size = Pt(9)
                        cr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

    doc.save(out_path)
    return out_path


# ---------- PDF (3-tier fallback) ----------

def _pdf_via_weasyprint(html_content: str, out_path: str) -> bool:
    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(out_path)
        return True
    except Exception as e:
        print(f"[pdf] WeasyPrint unavailable/failed: {e}")
        return False


def _pdf_via_wkhtmltopdf(html_content: str, out_path: str) -> bool:
    """Use wkhtmltopdf through pdfkit. Renders the styled HTML faithfully."""
    try:
        import pdfkit
        options = {
            "quiet": "",
            "encoding": "UTF-8",
            "enable-local-file-access": "",
            "margin-top": "16mm", "margin-bottom": "16mm",
            "margin-left": "14mm", "margin-right": "14mm",
        }
        pdfkit.from_string(html_content, out_path, options=options)
        return os.path.exists(out_path) and os.path.getsize(out_path) > 0
    except Exception as e:
        print(f"[pdf] wkhtmltopdf/pdfkit unavailable/failed: {e}")
        return False


def _pdf_via_reportlab(blocks: list, out_path: str, title: str) -> bool:
    """Last-resort PDF built directly from the block model with reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, Image, HRFlowable)
        from reportlab.lib.enums import TA_CENTER

        styles = getSampleStyleSheet()
        navy = colors.HexColor("#1f4e79")
        red = colors.HexColor("#c00000")

        h1 = ParagraphStyle("H1", parent=styles["Title"], textColor=navy,
                            fontSize=20, spaceAfter=6)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=navy,
                            fontSize=14, spaceBefore=16, spaceAfter=6)
        h3 = ParagraphStyle("H3", parent=styles["Heading3"], textColor=colors.HexColor("#2c5282"),
                            fontSize=11.5, spaceBefore=10, spaceAfter=4)
        body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9.5,
                              leading=14, spaceAfter=6)
        meta = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=8,
                              textColor=colors.HexColor("#777777"), spaceAfter=3)
        summ = ParagraphStyle("Summary", parent=styles["Normal"], fontSize=9.5,
                              leading=14, backColor=colors.HexColor("#f5f7fa"),
                              borderColor=navy, borderWidth=0, leftIndent=8,
                              spaceBefore=6, spaceAfter=8)
        cap = ParagraphStyle("Cap", parent=styles["Normal"], fontSize=8,
                             textColor=colors.HexColor("#777777"),
                             alignment=TA_CENTER, spaceAfter=8)

        def esc(t):
            return (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

        def inline(t):
            t = esc(t)
            t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
            t = re.sub(r"\*(.+?)\*", r"<i>\1</i>", t)
            return t

        story = []
        for b in blocks:
            t = b["type"]
            if t == "heading":
                if b["level"] == 1:
                    story.append(Paragraph(esc(b["text"]), h1))
                    story.append(HRFlowable(width="100%", thickness=2, color=red,
                                            spaceBefore=2, spaceAfter=8))
                elif b["level"] == 2:
                    story.append(Paragraph(esc(b["text"]), h2))
                else:
                    story.append(Paragraph(esc(b["text"]), h3))
            elif t == "paragraph":
                story.append(Paragraph(inline(b["text"]), body))
            elif t == "meta":
                story.append(Paragraph(inline(b["text"]), meta))
            elif t == "summary":
                story.append(Paragraph(inline(b["text"]), summ))
            elif t == "caption":
                story.append(Paragraph(inline(b["text"]), cap))
            elif t == "rule":
                story.append(HRFlowable(width="100%", thickness=0.5,
                                        color=colors.HexColor("#dddddd"),
                                        spaceBefore=10, spaceAfter=10))
            elif t == "table":
                headers = b["headers"]
                ncols = len(headers)
                data = [[Paragraph(f"<b>{esc(str(h))}</b>",
                                   ParagraphStyle("th", parent=body, fontSize=8.5,
                                                  textColor=colors.white))
                         for h in headers]]
                style_cmds = [
                    ("BACKGROUND", (0, 0), (-1, 0), navy),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
                ridx = 1
                for row in b["rows"]:
                    if isinstance(row, dict) and "_category" in row:
                        cell = Paragraph(f"<b>{esc(row['_category'])}</b>",
                                         ParagraphStyle("cat", parent=body,
                                                        fontSize=8.5, textColor=navy))
                        data.append([cell] + [""] * (ncols - 1))
                        style_cmds.append(("SPAN", (0, ridx), (-1, ridx)))
                        style_cmds.append(("BACKGROUND", (0, ridx), (-1, ridx),
                                           colors.HexColor("#eef2f7")))
                    else:
                        cells = [str(c) for c in row]
                        cells += [""] * (ncols - len(cells))
                        data.append([Paragraph(inline(c),
                                     ParagraphStyle("td", parent=body, fontSize=8))
                                     for c in cells])
                        if ridx % 2 == 0:
                            style_cmds.append(("BACKGROUND", (0, ridx), (-1, ridx),
                                               colors.HexColor("#fafbfc")))
                    ridx += 1
                avail = A4[0] - 28 * mm
                tbl = Table(data, colWidths=[avail / ncols] * ncols, repeatRows=1)
                tbl.setStyle(TableStyle(style_cmds))
                story.append(tbl)
                story.append(Spacer(1, 8))
            elif t == "chart":
                path = b.get("path", "")
                if path and os.path.exists(path):
                    try:
                        from PIL import Image as PILImage
                        iw, ih = PILImage.open(path).size
                        max_w = A4[0] - 50 * mm
                        w = min(max_w, iw * 0.55)
                        h = w * ih / iw
                        img = Image(path, width=w, height=h)
                        img.hAlign = "CENTER"
                        story.append(Spacer(1, 4))
                        story.append(img)
                        if b.get("caption"):
                            story.append(Paragraph(inline(b["caption"]), cap))
                        story.append(Spacer(1, 4))
                    except Exception as e:
                        print(f"[pdf] reportlab image skip: {e}")

        doc = SimpleDocTemplate(out_path, pagesize=A4,
                                topMargin=16 * mm, bottomMargin=16 * mm,
                                leftMargin=14 * mm, rightMargin=14 * mm,
                                title=title)
        doc.build(story)
        return os.path.exists(out_path) and os.path.getsize(out_path) > 0
    except Exception as e:
        print(f"[pdf] reportlab failed: {e}")
        return False


def _build_pdf(html_content: str, blocks: list, out_path: str,
               title: str) -> tuple[str | None, str]:
    """Try WeasyPrint -> wkhtmltopdf -> reportlab. Returns (path|None, engine)."""
    if _pdf_via_weasyprint(html_content, out_path):
        return out_path, "weasyprint"
    if _pdf_via_wkhtmltopdf(html_content, out_path):
        return out_path, "wkhtmltopdf"
    if _pdf_via_reportlab(blocks, out_path, title):
        return out_path, "reportlab"
    return None, "none"


# =========================================================================
# Agent entry point
# =========================================================================

def report_writer_agent(state: AnalysisState) -> AnalysisState:
    log = state.get("log", [])
    working_dir = state.get("working_dir", "/tmp")
    tickers = state.get("tickers", [])

    title = (f"{tickers[0]} Financial Analysis" if len(tickers) == 1
             else "Financial Statement Analysis Report")

    blocks = _build_blocks(state)
    md_text = _build_markdown(blocks)
    html_text = _build_html(blocks, title)

    docx_path = os.path.join(working_dir, "financial_analysis.docx")
    pdf_path = os.path.join(working_dir, "financial_analysis.pdf")

    try:
        _build_docx(blocks, docx_path)
        log.append(f"Report Writer: DOCX written -> {docx_path}")
    except Exception as e:
        docx_path = ""
        log.append(f"Report Writer: DOCX failed ({e})")

    pdf_result, engine = _build_pdf(html_text, blocks, pdf_path, title)
    if pdf_result:
        pdf_path = pdf_result
        log.append(f"Report Writer: PDF written via {engine} -> {pdf_path}")
    else:
        pdf_path = ""
        log.append("Report Writer: PDF generation failed on all engines")

    log.append(f"Report Writer: {len(blocks)} blocks, markdown {len(md_text)} chars, "
               f"HTML {len(html_text)} chars")

    return {**state,
            "final_report_md": md_text,
            "final_report_html": html_text,
            "final_report_docx_path": docx_path,
            "final_report_pdf_path": pdf_path,
            "log": log}
